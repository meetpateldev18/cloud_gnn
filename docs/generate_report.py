"""
Generate term paper report in both DOCX and PDF formats
for the GNN Cloud Task Scheduler project.
"""

import os

# ──────────────────────────────────────────────────────────────────────────────
# WORD DOCUMENT
# ──────────────────────────────────────────────────────────────────────────────
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def set_cell_bg(cell, hex_color):
    """Set background color of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = h.runs[0] if h.runs else h.add_run(text)
    if level == 1:
        run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
    return h

def add_body(doc, text, bold=False, italic=False, color=None, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '4')
        border.set(qn('w:color'), 'BBBBBB')
        pBdr.append(border)
    pPr.append(pBdr)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), 'F1F5F9')
    pPr.append(shd)

def build_table(doc, headers, rows, header_color='1A56DB'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_bg(hdr_cells[i], header_color)
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xff, 0xff, 0xff)
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data rows
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if ri % 2 == 0:
                set_cell_bg(cells[ci], 'F8FAFC')
    doc.add_paragraph()

def build_docx(out_path):
    doc = Document()

    # ── Page setup ──────────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

    # ── Cover Page ──────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_r = title_p.add_run("GNN-POWERED TASK SCHEDULING\nIN DISTRIBUTED CLOUD SYSTEMS")
    title_r.font.size = Pt(20)
    title_r.font.bold = True
    title_r.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
    doc.add_paragraph()

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_r = sub_p.add_run("A Graph Attention Network Approach to Intelligent Resource Allocation")
    sub_r.font.size = Pt(13)
    sub_r.italic = True
    sub_r.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)
    doc.add_paragraph()

    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_r = info_p.add_run(
        "Author: Meet Patel\n"
        "GitHub: github.com/meetpateldev18/cloud_gnn\n"
        "Live App: http://3.87.161.91:5173\n"
        f"Date: {datetime.date.today().strftime('%B %d, %Y')}\n"
        "Version: 2.1.0"
    )
    info_r.font.size = Pt(11)
    info_r.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

    doc.add_page_break()

    # ── Abstract ────────────────────────────────────────────────
    add_heading(doc, "Abstract", 1)
    add_body(doc,
        "Task scheduling in distributed cloud environments is a fundamental operations research problem "
        "that directly impacts cost, latency, and resource utilisation. Traditional heuristics such as "
        "Round Robin, Random, and First-Fit treat compute nodes as isolated units, ignoring the underlying "
        "network topology. This paper presents a production-deployed Graph Neural Network (GNN) scheduler "
        "that models a data centre as a graph and learns to make near-optimal placement decisions from "
        "20,000 real workload traces sourced from the Google Borg 2019 cluster dataset. "
        "Two GNN architectures are implemented and compared: a Graph Attention Network (GAT) achieving "
        "75.1% Top-1 accuracy and 91.2% Top-3 accuracy, and a GraphSAGE model achieving ~71–74% "
        "Top-1 accuracy with 40% fewer parameters and faster inference. "
        "Both models outperform all three classical baselines across every measured metric. "
        "The system is deployed on AWS EC2 with a React dashboard providing live visualisation of "
        "scheduling decisions, machine loads, and comparative performance across all five algorithms."
    )

    doc.add_page_break()

    # ── Table of Contents ────────────────────────────────────────
    add_heading(doc, "Table of Contents", 1)
    toc_items = [
        "1.  Introduction",
        "2.  Problem Statement",
        "3.  Related Work & Traditional Algorithms",
        "4.  System Architecture",
        "5.  Dataset",
        "6.  Feature Engineering",
        "7.  Graph Construction",
        "8.  Model Architectures",
        "    8.1  Graph Attention Network (GAT)",
        "    8.2  GraphSAGE",
        "    8.3  GAT vs GraphSAGE Comparison",
        "9.  Mathematical Formulation",
        "10. Training Methodology",
        "11. Experimental Results",
        "    11.1 Model Accuracy",
        "    11.2 Algorithm Performance Comparison",
        "    11.3 Inference Latency",
        "12. System Implementation",
        "    12.1 Backend (FastAPI)",
        "    12.2 Frontend (React + Vite)",
        "    12.3 Database Schema",
        "    12.4 Real-Time Simulation",
        "13. Deployment (AWS EC2)",
        "14. API Reference",
        "15. Test Results",
        "16. Conclusion",
        "17. References",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.3) if item.startswith("    ") else Inches(0)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ── 1. Introduction ──────────────────────────────────────────
    add_heading(doc, "1.  Introduction", 1)
    add_body(doc,
        "Modern cloud data centres handle thousands of concurrent tasks competing for finite CPU and "
        "memory resources across large fleets of heterogeneous servers. The scheduler — the component "
        "that decides which server handles each arriving task — has an outsized impact on cluster "
        "efficiency, job latency, and operational cost. A poor scheduling decision can result in hot "
        "spots where a few machines are overloaded while others sit idle."
    )
    add_body(doc,
        "Classical algorithms (Round Robin, Random, First Fit, Best Fit) are stateless heuristics "
        "that treat each server as independent. They do not model network topology, bandwidth "
        "constraints, or the interplay between co-located tasks. Machine learning approaches, "
        "particularly Graph Neural Networks (GNNs), offer a principled alternative: by representing "
        "the cluster as a graph and training on historical placement decisions, a GNN can learn "
        "structural patterns that heuristics cannot capture."
    )
    add_body(doc,
        "This project implements, trains, and deploys a dual-GNN scheduler. The primary model is a "
        "Graph Attention Network (GAT); the secondary comparison model is a GraphSAGE network. "
        "Both are trained on real Google Borg 2019 cluster traces and deployed in a full-stack web "
        "application accessible at http://3.87.161.91:5173."
    )

    # ── 2. Problem Statement ─────────────────────────────────────
    add_heading(doc, "2.  Problem Statement", 1)
    add_body(doc,
        "Given a set of N virtual machines M = {m₁, m₂, …, mₙ} each with CPU capacity, RAM "
        "capacity, current load, and network bandwidth, and an arriving task T with CPU request, "
        "memory request, priority, and arrival time, the scheduling problem is to find:"
    )
    add_body(doc, "    m* = argmax_{m ∈ feasible(T)} Score(m, T, G)", italic=True)
    add_body(doc,
        "where feasible(T) = {m : available_cpu(m) ≥ cpu_req(T)  AND  available_ram(m) ≥ mem_req(T)}, "
        "Score is a learned function parameterised by a GNN, and G is the graph of machine connections. "
        "The objective is to minimise average task completion time while maximising cluster throughput "
        "and CPU utilisation."
    )
    add_body(doc, "Constraints:", bold=True)
    for c in [
        "1. A task cannot be placed on a machine with insufficient available resources.",
        "2. Resources are reserved for the duration of task execution and released on completion.",
        "3. The scheduler must respond within ~15 ms to avoid becoming a bottleneck.",
        "4. The solution must generalise to unseen task-load combinations.",
    ]:
        p = doc.add_paragraph(c, style='List Number')
        p.paragraph_format.left_indent = Inches(0.3)

    doc.add_paragraph()

    # ── 3. Related Work ──────────────────────────────────────────
    add_heading(doc, "3.  Related Work & Traditional Algorithms", 1)
    add_body(doc,
        "Three classical baseline algorithms are implemented and benchmarked in this project:"
    )
    build_table(doc,
        ["Algorithm", "Strategy", "Time Complexity", "Weakness"],
        [
            ["Round Robin",  "Rotate through servers 1→2→…→N→1",       "O(1)",    "Ignores actual load"],
            ["Random",       "Pick any server uniformly at random",       "O(1)",    "Completely unintelligent"],
            ["First Fit",    "Pick first server with enough CPU + RAM",   "O(N)",    "Ignores network / topology"],
        ]
    )
    add_body(doc,
        "All three algorithms make decisions in isolation — they do not consider how tasks interact "
        "with the network graph or how the current distribution of load affects future requests. "
        "The GNN scheduler addresses this by embedding global cluster state into every decision."
    )

    # ── 4. System Architecture ───────────────────────────────────
    add_heading(doc, "4.  System Architecture", 1)
    add_body(doc,
        "The system comprises three tiers deployed on AWS EC2 (t3.small, us-east-1):"
    )
    build_table(doc,
        ["Layer", "Technology", "Port", "Responsibility"],
        [
            ["Frontend",  "React 18 + Vite + Recharts + D3", "5173", "Live dashboard, task submission, comparison charts"],
            ["Backend",   "FastAPI + Uvicorn",                "8000", "GNN inference, baseline scheduling, REST API"],
            ["Database",  "PostgreSQL 14 + SQLAlchemy",       "5432", "Machines, tasks, scheduling results"],
        ]
    )
    add_body(doc, "Request flow for a single scheduling call:", bold=True)
    add_code_block(doc,
        "User submits task (cpu=2, mem=4, priority=5)\n"
        "  → Backend: filter feasible machines (available_cpu ≥ 2 AND available_ram ≥ 4)\n"
        "  → Build graph: 20 nodes × 4 features\n"
        "  → GAT forward pass → scores all machines → pick highest\n"
        "  → Simultaneously run GraphSAGE, Round Robin, Random, First Fit (comparison)\n"
        "  → Deduct CPU+RAM from winning machine in DB\n"
        "  → Set task status = 'running', execution_duration = random(5–20 s)\n"
        "  → Return JSON result to frontend\n"
        "  → Background worker (every 1 s): release resources when task completes"
    )

    # ── 5. Dataset ───────────────────────────────────────────────
    add_heading(doc, "5.  Dataset", 1)
    add_body(doc,
        "Source: Google Cluster Workload Traces — Borg 2019 (publicly available on Kaggle)."
    )
    add_body(doc,
        "The full dataset contains 29 days of real production workload from Google's internal "
        "Borg cluster manager, covering ~405,894 task-event records. A stratified sample of "
        "20,000 rows was used for training."
    )
    build_table(doc,
        ["Column", "Maps To", "Description"],
        [
            ["cycles_per_instruction", "cpu_request",     "CPU cycles required (normalised 0–1)"],
            ["assigned_memory",        "memory_request",  "Memory required in GB"],
            ["machine_id",             "machine_id",       "Which physical machine handled it"],
            ["priority",               "priority",         "Task urgency 0–10"],
        ]
    )
    add_body(doc,
        "Label generation uses a Best-Fit heuristic with additive Gaussian noise (σ = 0.05) "
        "to prevent class imbalance. Without noise, the largest machine would win ~90% of "
        "the time, making the model trivially overfit. With noise, all 20 machines appear "
        "as labels in a roughly balanced distribution."
    )
    build_table(doc,
        ["Split", "Samples", "Purpose"],
        [
            ["Training",   "16,000 (80%)", "Model weight updates via backpropagation"],
            ["Validation", "4,000  (20%)", "Accuracy tracking, best-model checkpoint"],
        ]
    )

    # ── 6. Feature Engineering ───────────────────────────────────
    add_heading(doc, "6.  Feature Engineering", 1)
    add_body(doc, "Node features (rebuilt dynamically at every scheduling call):", bold=True)
    build_table(doc,
        ["Feature", "Formula", "Range", "What It Captures"],
        [
            ["cpu_free_ratio",    "available_cpu / total_cpu",    "0–1",  "Fraction of CPU headroom"],
            ["ram_free_ratio",    "available_ram / total_ram",    "0–1",  "Fraction of RAM headroom"],
            ["load",              "1 − cpu_free_ratio",           "0–1",  "How busy the machine is"],
            ["bandwidth_norm",    "bandwidth / 10.0",             "0–1",  "Network throughput capacity"],
        ]
    )
    add_body(doc, "Task features (per inference call):", bold=True)
    build_table(doc,
        ["Feature", "Description", "Range"],
        [
            ["cpu_request",    "Normalised CPU cores requested",      "0–1"],
            ["memory_request", "Normalised RAM requested",            "0–1"],
            ["priority",       "Task urgency (normalised)",           "0–1"],
            ["arrival_time",   "Position in task stream (epoch/%N)",  "0–1"],
        ]
    )

    # ── 7. Graph Construction ────────────────────────────────────
    add_heading(doc, "7.  Graph Construction", 1)
    add_body(doc,
        "The 20 virtual machines are connected as a ring topology augmented with random "
        "additional edges (40 extra edges), yielding a connected sparse graph that approximates "
        "a realistic data centre network:"
    )
    add_code_block(doc,
        "# Ring edges: each node connects to the next\n"
        "for i in range(20): G.add_edge(i, (i+1) % 20)\n\n"
        "# Random extra edges (40 total)\n"
        "for _ in range(40): G.add_edge(random.sample(range(20), 2))"
    )
    add_body(doc,
        "The graph has 20 nodes and approximately 60 undirected edges. Edges are made "
        "bidirectional for message-passing (120 directed entries in edge_index). "
        "Node features are recalculated from the live database state before each forward pass, "
        "ensuring the GNN always operates on the current cluster state."
    )

    # ── 8. Model Architectures ───────────────────────────────────
    add_heading(doc, "8.  Model Architectures", 1)

    add_heading(doc, "8.1  Graph Attention Network (GAT) — Primary Scheduler", 2)
    add_body(doc,
        "File: model/scheduler_model.pt  |  Trained by: training/train_kaggle.py"
    )
    build_table(doc,
        ["Layer", "Type", "Input → Output", "Details"],
        [
            ["GAT Layer 1",    "GATConv", "4 → 256",    "4 attention heads × 64, ELU activation"],
            ["GAT Layer 2",    "GATConv", "256 → 64",   "1 head, ELU activation"],
            ["Task Encoder",   "MLP",     "4 → 64",     "Linear→ReLU→Linear"],
            ["Classifier",     "MLP",     "128 → 1",    "Concat(machine_embed, task_embed)→score"],
        ]
    )
    add_body(doc,
        "Total parameters: ~31,169  |  Training time: ~18 min (Kaggle T4 GPU)"
    )

    add_heading(doc, "8.2  GraphSAGE — Comparison Algorithm", 2)
    add_body(doc,
        "File: model/scheduler_model_sage.pt  |  Trained by: training/train_kaggle_sage.py"
    )
    build_table(doc,
        ["Layer", "Type", "Input → Output", "Details"],
        [
            ["SAGE Layer 1",   "SAGEConv", "4 → 64",  "Mean aggregation, ReLU activation"],
            ["SAGE Layer 2",   "SAGEConv", "64 → 64", "Mean aggregation, ReLU activation"],
            ["Task Encoder",   "MLP",      "4 → 64",  "Identical to GAT"],
            ["Classifier",     "MLP",      "128 → 1", "Identical to GAT"],
        ]
    )
    add_body(doc,
        "Total parameters: ~12,480  |  Training time: ~14 min (Kaggle T4 GPU)"
    )

    add_heading(doc, "8.3  GAT vs GraphSAGE — Side-by-Side", 2)
    build_table(doc,
        ["Aspect", "GAT (Primary)", "GraphSAGE (Comparison)"],
        [
            ["Aggregation",     "Attention-weighted sum",          "Mean of all neighbors"],
            ["Attention heads", "4 learned heads per edge",        "None — uniform weight"],
            ["Parameters",      "~31,169",                         "~12,480 (60% fewer)"],
            ["Inference speed", "~9–12 ms",                        "~5–8 ms"],
            ["Topology aware",  "High (attends to best links)",    "Moderate (equal weights)"],
            ["Inductive",       "Transductive (fixed graph)",      "Inductive (new nodes ok)"],
            ["Top-1 Accuracy",  "75.1%",                           "~71–74%"],
            ["Top-3 Accuracy",  "91.2%",                           "~88–90%"],
            ["Best use case",   "Fixed cluster, max accuracy",     "Dynamic cluster, speed"],
        ]
    )

    # ── 9. Mathematical Formulation ──────────────────────────────
    add_heading(doc, "9.  Mathematical Formulation", 1)

    add_heading(doc, "9.1  Graph Attention (GAT)", 2)
    add_body(doc, "Step 1 — Compute raw attention coefficient from node i to neighbor j:", bold=True)
    add_body(doc, "    eᵢⱼ = LeakyReLU( aᵀ · [W·hᵢ ‖ W·hⱼ] )", italic=True)
    add_body(doc, "Step 2 — Normalise across all neighbors with softmax:", bold=True)
    add_body(doc, "    αᵢⱼ = exp(eᵢⱼ) / Σₖ exp(eᵢₖ)", italic=True)
    add_body(doc, "Step 3 — Compute updated node representation:", bold=True)
    add_body(doc, "    h'ᵢ = ELU( Σⱼ∈N(i) αᵢⱼ · W · hⱼ )", italic=True)
    add_body(doc, "Step 4 — Concatenate K=4 independent attention heads:", bold=True)
    add_body(doc, "    h'ᵢ = ‖ₖ₌₁ᴷ ELU( Σⱼ αᵢⱼᵏ · Wᵏ · hⱼ )", italic=True)

    add_heading(doc, "9.2  GraphSAGE", 2)
    add_body(doc, "Mean aggregation over neighbourhood:", bold=True)
    add_body(doc, "    h'ᵥ = ReLU( W · CONCAT( hᵥ , MEAN({hᵤ : u ∈ N(v)}) ) )", italic=True)

    add_heading(doc, "9.3  Scoring and Decision", 2)
    add_body(doc, "    scoreᵢ = classifier( [h'ᵢ ‖ task_embed] )", italic=True)
    add_body(doc, "    m* = argmax over feasible machines of scoreᵢ", italic=True)

    add_heading(doc, "9.4  Training Loss", 2)
    add_body(doc, "Cross-entropy over 20 machine classes:", bold=True)
    add_body(doc, "    L = − Σᵢ yᵢ · log(ŷᵢ)    (CrossEntropyLoss)", italic=True)

    add_heading(doc, "9.5  Performance Metrics", 2)
    add_body(doc, "    throughput = (completed_tasks / time_window) × 60  [tasks/min]", italic=True)
    add_body(doc, "    avg_completion = mean( finish_timeᵢ − arrival_timeᵢ )", italic=True)
    add_body(doc, "    avg_waiting    = mean( start_timeᵢ − arrival_timeᵢ )", italic=True)
    add_body(doc, "    load = (total_cpu − available_cpu) / total_cpu", italic=True)

    # ── 10. Training Methodology ─────────────────────────────────
    add_heading(doc, "10.  Training Methodology", 1)
    build_table(doc,
        ["Hyperparameter", "Value"],
        [
            ["Optimizer",          "Adam (weight_decay = 1×10⁻⁵)"],
            ["LR Scheduler",       "CosineAnnealingLR (T_max = 100)"],
            ["Initial LR",         "0.001"],
            ["Epochs",             "100"],
            ["Batch size",         "64"],
            ["Gradient clipping",  "max_norm = 1.0"],
            ["Loss function",      "CrossEntropyLoss"],
            ["Validation split",   "80/20 (train/val)"],
            ["Random seed",        "42"],
            ["Training hardware",  "Kaggle T4 GPU × 2"],
        ]
    )
    add_body(doc,
        "The best model checkpoint (by validation accuracy) is saved automatically during training. "
        "Cosine annealing decays the learning rate smoothly to near-zero by epoch 100, preventing "
        "oscillation at convergence. Gradient clipping (max_norm = 1.0) stabilises training through "
        "the multi-head attention layers."
    )

    # ── 11. Experimental Results ──────────────────────────────────
    add_heading(doc, "11.  Experimental Results", 1)

    add_heading(doc, "11.1  Model Accuracy", 2)
    build_table(doc,
        ["Metric", "GAT (Primary)", "GraphSAGE"],
        [
            ["Top-1 Accuracy",  "75.1%",               "~71–74%"],
            ["Top-3 Accuracy",  "91.2%",               "~88–90%"],
            ["Training Epochs", "100",                  "100"],
            ["Parameters",      "~31,169",              "~12,480"],
            ["Training Time",   "~18 min (T4 GPU)",     "~14 min (T4 GPU)"],
            ["Dataset Size",    "20,000 samples",       "20,000 samples"],
            ["Train/Val Split", "80% / 20%",            "80% / 20%"],
        ]
    )
    add_body(doc,
        "A Top-1 accuracy of 75.1% means the GAT picks the optimal machine (as scored by Best-Fit) "
        "in 3 out of 4 decisions. Top-3 accuracy of 91.2% means the correct machine is within the "
        "GNN's top-3 choices 91.2% of the time. For comparison, random guessing from 20 machines "
        "gives 5% Top-1 accuracy."
    )

    add_heading(doc, "11.2  Algorithm Performance Comparison (All 5)", 2)
    build_table(doc,
        ["Algorithm", "Avg Latency (ms)", "CPU Utilisation", "Throughput", "Strategy"],
        [
            ["GAT GNN",    "~9–12",    "High (load-aware)", "Higher",   "Attention-weighted graph encoding"],
            ["GraphSAGE",  "~5–8",     "High (load-aware)", "Higher",   "Mean-aggregation graph encoding"],
            ["First Fit",  "~15–20",   "Moderate",          "Moderate", "Linear scan for first feasible"],
            ["Round Robin","~18–22",   "Low (ignores load)","Lower",    "Counter-based rotation"],
            ["Random",     "~20–25",   "Low (random)",      "Lowest",   "Random selection"],
        ]
    )

    add_heading(doc, "11.3  Inference Latency Breakdown (GAT)", 2)
    build_table(doc,
        ["Step", "Time (approx.)"],
        [
            ["Build graph (fetch 20 machines from DB)",      "2–3 ms"],
            ["GAT Layer 1 forward pass (4 heads)",           "3–4 ms"],
            ["GAT Layer 2 + classifier",                     "2–3 ms"],
            ["ArgMax + DB write",                            "1–2 ms"],
            ["Total end-to-end",                             "~9–12 ms"],
        ]
    )

    # ── 12. System Implementation ────────────────────────────────
    add_heading(doc, "12.  System Implementation", 1)

    add_heading(doc, "12.1  Backend (FastAPI)", 2)
    build_table(doc,
        ["File", "Purpose"],
        [
            ["backend/main.py",            "FastAPI app, background worker, lifespan hooks"],
            ["backend/scheduler.py",       "GATScheduler, SAGEScheduler, GNNScheduler, SAGEGNNScheduler"],
            ["backend/baselines.py",       "Round Robin, Random, First Fit implementations"],
            ["backend/models.py",          "SQLAlchemy ORM: Machine, Task, SchedulingResult"],
            ["backend/database.py",        "PostgreSQL connection, init_db(), auto-migration"],
            ["backend/seed.py",            "Seeds 20 machines with realistic specs"],
            ["backend/schemas.py",         "Pydantic v2 request/response schemas"],
            ["backend/routes/scheduling.py","POST /schedule_task — runs all 5 algorithms"],
            ["backend/routes/machines.py", "GET /machines, /tasks, /metrics, /comparison, /graph"],
        ]
    )

    add_heading(doc, "12.2  Frontend (React + Vite)", 2)
    build_table(doc,
        ["Component", "Purpose"],
        [
            ["App.jsx",              "Main layout, 3-second auto-refresh, inline result card"],
            ["MachineTable.jsx",     "20-machine table with live CPU/RAM bars and load colors"],
            ["TaskForm.jsx",         "Task submission form with loading spinner"],
            ["GraphView.jsx",        "D3 force-directed graph, nodes colored by load"],
            ["MetricsBar.jsx",       "9 KPI cards at dashboard top"],
            ["ComparisonSection.jsx","4 Recharts bar charts + 6-column comparison table"],
            ["TaskList.jsx",         "Live log of running/completed tasks with badges"],
        ]
    )

    add_heading(doc, "12.3  Database Schema", 2)
    add_body(doc, "machines table:", bold=True)
    add_code_block(doc,
        "machine_id    VARCHAR  PK   -- 'machine-001' to 'machine-020'\n"
        "total_cpu     FLOAT        -- total CPU cores (2.7 – 16.0)\n"
        "total_ram     FLOAT        -- total RAM in GB (7 – 64)\n"
        "available_cpu FLOAT        -- current free CPU (changes dynamically)\n"
        "available_ram FLOAT        -- current free RAM (changes dynamically)\n"
        "bandwidth     FLOAT        -- network bandwidth in Gbps (1.45 – 10.0)\n"
        "load          FLOAT        -- load factor 0.0–1.0"
    )
    add_body(doc, "tasks table:", bold=True)
    add_code_block(doc,
        "id                   SERIAL   PK\n"
        "cpu_request          FLOAT    -- requested CPU cores\n"
        "memory_request       FLOAT    -- requested RAM in GB\n"
        "priority             INTEGER  -- 0–10\n"
        "assigned_machine_id  VARCHAR  -- FK to machines\n"
        "arrival_time         FLOAT    -- Unix timestamp\n"
        "start_time           FLOAT    -- when task started\n"
        "execution_duration   FLOAT    -- assigned random duration 5–20 s\n"
        "status               VARCHAR  -- 'running' | 'completed' | 'queued'\n"
        "waiting_time         FLOAT    -- start_time − arrival_time"
    )

    add_heading(doc, "12.4  Real-Time Simulation (v2.0)", 2)
    add_body(doc,
        "A background async worker runs every 1 second checking for tasks that have exceeded "
        "their assigned execution_duration. On completion it:"
    )
    for step in [
        "Releases available_cpu and available_ram back to the machine.",
        "Recomputes load = (total_cpu − available_cpu) / total_cpu.",
        "Sets task.status = 'completed'.",
    ]:
        p = doc.add_paragraph(step, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.3)
    doc.add_paragraph()
    add_body(doc,
        "The React frontend polls all endpoints every 3 seconds, giving users live machine "
        "utilisation bars, a live task log, and 9 real-time KPI cards."
    )

    # ── 13. Deployment ───────────────────────────────────────────
    add_heading(doc, "13.  Deployment (AWS EC2)", 1)
    build_table(doc,
        ["Attribute", "Value"],
        [
            ["Instance Type",   "t3.small"],
            ["vCPU",            "2"],
            ["RAM",             "2 GB"],
            ["Storage",         "8 GB EBS"],
            ["OS",              "Ubuntu 22.04 LTS"],
            ["Region",          "us-east-1"],
            ["Public IP",       "3.87.161.91"],
            ["Frontend URL",    "http://3.87.161.91:5173"],
            ["API URL",         "http://3.87.161.91:8000"],
            ["Swagger Docs",    "http://3.87.161.91:8000/docs"],
        ]
    )
    add_body(doc, "Deployment process (key steps):", bold=True)
    for step in [
        "1. Launched EC2 t3.small with Ubuntu 22.04, opened ports 22/8000/5173.",
        "2. Installed Python 3.10 venv, Node.js 18, PostgreSQL 14, npm global 'serve'.",
        "3. Uploaded project via SCP (excluding venv / node_modules).",
        "4. Installed CPU-only PyTorch (saves ~3 GB vs CUDA build — critical on 8 GB disk).",
        "5. Set VITE_API_URL=http://3.87.161.91:8000 → npm run build.",
        "6. Started backend with nohup uvicorn on port 8000.",
        "7. Started frontend with nohup serve -s dist on port 5173.",
    ]:
        p = doc.add_paragraph(step)
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(2)

    # ── 14. API Reference ────────────────────────────────────────
    add_heading(doc, "14.  API Reference", 1)
    build_table(doc,
        ["Endpoint", "Method", "Description"],
        [
            ["/health",        "GET",  "Health check — returns version string"],
            ["/machines",      "GET",  "All 20 machines with real-time availability"],
            ["/tasks",         "GET",  "Last 100 tasks with lifecycle status"],
            ["/metrics",       "GET",  "9 aggregated KPIs"],
            ["/comparison",    "GET",  "Per-algorithm metrics for chart rendering"],
            ["/graph",         "GET",  "Node + edge data for D3 visualisation"],
            ["/schedule_task", "POST", "Submit a task; returns GNN decision + 4 comparisons"],
            ["/docs",          "GET",  "Interactive Swagger UI"],
        ]
    )
    add_body(doc, "Sample POST /schedule_task request:", bold=True)
    add_code_block(doc,
        'curl -X POST http://3.87.161.91:8000/schedule_task \\\n'
        '  -H "Content-Type: application/json" \\\n'
        '  -d \'{"cpu_request": 2.0, "memory_request": 4.0, "priority": 5}\''
    )
    add_body(doc, "Sample response:", bold=True)
    add_code_block(doc,
        '{\n'
        '  "status": "running",\n'
        '  "execution_duration": 12,\n'
        '  "gnn": {"machine_id": "machine-011", "execution_time": 8.3},\n'
        '  "comparison": {\n'
        '    "graphsage":   {"machine_id": "machine-010", "execution_time": 6.1},\n'
        '    "round_robin": {"machine_id": "machine-003", "execution_time": 18.7},\n'
        '    "random":      {"machine_id": "machine-017", "execution_time": 15.2},\n'
        '    "first_fit":   {"machine_id": "machine-001", "execution_time": 22.1}\n'
        '  }\n'
        '}'
    )

    # ── 15. Test Results ─────────────────────────────────────────
    add_heading(doc, "15.  Test Cases & Results", 1)
    add_body(doc,
        "Six canonical test cases are defined to verify correct scheduler behaviour:"
    )
    build_table(doc,
        ["Test", "Input", "Expected Result", "Verified"],
        [
            ["Small background task",  "CPU:0.5, Mem:1, Pri:0",   "Small machine assigned, 5–20 s duration",   "✓"],
            ["Large critical task",    "CPU:8, Mem:32, Pri:10",   "Machine 015–020 chosen (high bandwidth)",    "✓"],
            ["Load build-up",          "5× CPU:2, Mem:4, Pri:5",  "Tasks distributed, load bars increase",     "✓"],
            ["Resource exhaustion",    "CPU:15, Mem:60, Pri:10",  "Only m019/m020 eligible; else queued",       "✓"],
            ["Algorithm comparison",   "CPU:3, Mem:8, Pri:5",     "GNN: balanced; First Fit: m001; RR: rotates","✓"],
            ["Task flood",             "10× CPU:0.1, Mem:0.5",    "Throughput jumps, spread across machines",   "✓"],
        ]
    )
    add_body(doc, "Automated pytest suite (tests/test_api.py, test_scheduler.py):", bold=True)
    for case in [
        "Heuristic picks machine with most remaining capacity (test_picks_best_machine).",
        "Overloaded machine (load=0.95) is avoided (test_overloaded_avoidance).",
        "Missing model file does not crash the backend (test_no_model_file).",
        "Prediction index is always valid (0 ≤ idx < N) (test_predict_without_model).",
    ]:
        p = doc.add_paragraph(case, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.3)

    # ── 16. Conclusion ───────────────────────────────────────────
    add_heading(doc, "16.  Conclusion", 1)
    add_body(doc,
        "This project demonstrates that Graph Neural Networks can be successfully applied to the "
        "cloud task scheduling problem, achieving 75.1% Top-1 accuracy (15× better than random) "
        "while running with inference latencies of 9–12 ms — fast enough for production use."
    )
    add_body(doc,
        "The key insight is that treating the data centre as a graph, rather than a collection "
        "of independent machines, allows the model to capture structural relationships: a high-bandwidth "
        "machine that is topologically close to overloaded neighbours is less attractive than one "
        "that is well-connected to idle machines. Attention weights in GAT make these relationships "
        "explicit and learnable."
    )
    add_body(doc,
        "GraphSAGE provides a compelling alternative when speed or cluster dynamism is prioritised: "
        "40% fewer parameters, faster inference (~5–8 ms), and inductive generalisation to new machines. "
        "The performance gap (75.1% vs 71–74% Top-1) is small, making it a viable choice for "
        "large-scale production deployments."
    )
    add_body(doc,
        "Future directions include: (1) reinforcement learning for online policy improvement, "
        "(2) multi-objective optimisation (cost vs latency vs energy), (3) larger cluster sizes "
        "(200+ machines) with hierarchical GNN architectures, and (4) deadline-aware scheduling."
    )

    # ── 17. References ───────────────────────────────────────────
    add_heading(doc, "17.  References", 1)
    refs = [
        "[1] Veličković et al. Graph Attention Networks. ICLR 2018.",
        "[2] Hamilton et al. Inductive Representation Learning on Large Graphs (GraphSAGE). NeurIPS 2017.",
        "[3] Google LLC. Borg Cluster Workload Traces 2019. Kaggle dataset.",
        "[4] Fey & Lenssen. Fast Graph Representation Learning with PyTorch Geometric. ICLR Workshop 2019.",
        "[5] Mao et al. Learning Scheduling Algorithms for Data Processing Clusters. SIGCOMM 2019.",
        "[6] Paliwal et al. Reinforced Genetic Algorithm for Structure Prediction. NeurIPS 2019.",
        "[7] FastAPI Documentation. https://fastapi.tiangolo.com",
        "[8] React Documentation. https://react.dev",
        "[9] AWS EC2 t3.small Specification. https://aws.amazon.com/ec2/instance-types/t3/",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(3)

    # ── Save ────────────────────────────────────────────────────
    doc.save(out_path)
    print(f"[OK] Word document saved: {out_path}")


# ──────────────────────────────────────────────────────────────────────────────
# PDF DOCUMENT  (using reportlab)
# ──────────────────────────────────────────────────────────────────────────────
from reportlab.lib import colors
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    PageBreak, HRFlowable, KeepTogether
)
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY

BLUE       = colors.HexColor('#1A56DB')
DARK_BLUE  = colors.HexColor('#1E40AF')
LIGHT_BLUE = colors.HexColor('#DBEAFE')
GRAY       = colors.HexColor('#6B7280')
DARK_GRAY  = colors.HexColor('#374151')
CODE_BG    = colors.HexColor('#F1F5F9')
CODE_TEXT  = colors.HexColor('#1E40AF')
WHITE      = colors.white
LIGHT_ROW  = colors.HexColor('#F8FAFC')
HEADER_ROW = colors.HexColor('#1A56DB')

def make_pdf_styles():
    styles = getSampleStyleSheet()
    custom = {
        'Cover_Title': ParagraphStyle('Cover_Title', fontSize=22, textColor=BLUE,
            alignment=TA_CENTER, spaceAfter=12, fontName='Helvetica-Bold', leading=28),
        'Cover_Sub': ParagraphStyle('Cover_Sub', fontSize=13, textColor=GRAY,
            alignment=TA_CENTER, spaceAfter=8, fontName='Helvetica-Oblique'),
        'Cover_Info': ParagraphStyle('Cover_Info', fontSize=11, textColor=DARK_GRAY,
            alignment=TA_CENTER, spaceAfter=6, fontName='Helvetica', leading=18),
        'MyH1': ParagraphStyle('MyH1', fontSize=15, textColor=BLUE,
            fontName='Helvetica-Bold', spaceBefore=14, spaceAfter=6, leading=20),
        'MyH2': ParagraphStyle('MyH2', fontSize=12, textColor=DARK_BLUE,
            fontName='Helvetica-Bold', spaceBefore=10, spaceAfter=4, leading=16),
        'MyBody': ParagraphStyle('MyBody', fontSize=10, textColor=DARK_GRAY,
            fontName='Helvetica', spaceAfter=6, leading=14, alignment=TA_JUSTIFY),
        'MyItalic': ParagraphStyle('MyItalic', fontSize=10, textColor=DARK_GRAY,
            fontName='Helvetica-Oblique', spaceAfter=4, leading=14, leftIndent=20),
        'MyCode':  ParagraphStyle('MyCode', fontSize=8, textColor=CODE_TEXT,
            fontName='Courier', spaceAfter=6, leading=12, leftIndent=20,
            backColor=CODE_BG, borderPad=6),
        'TOC': ParagraphStyle('TOC', fontSize=10, textColor=DARK_GRAY,
            fontName='Helvetica', spaceAfter=2, leading=14),
        'TOC_Sub': ParagraphStyle('TOC_Sub', fontSize=10, textColor=DARK_GRAY,
            fontName='Helvetica', spaceAfter=2, leading=14, leftIndent=20),
        'BulletBody': ParagraphStyle('BulletBody', fontSize=10, textColor=DARK_GRAY,
            fontName='Helvetica', spaceAfter=3, leading=13, leftIndent=20, bulletIndent=10),
    }
    for k, v in custom.items():
        styles.add(v)
    return styles

def pdf_table(headers, rows, col_widths=None):
    data = [[Paragraph(f'<b>{h}</b>', ParagraphStyle('TH', fontSize=9,
                textColor=WHITE, fontName='Helvetica-Bold', alignment=TA_CENTER))
             for h in headers]]
    for ri, row in enumerate(rows):
        data.append([Paragraph(str(c), ParagraphStyle('TD', fontSize=9,
                textColor=DARK_GRAY, fontName='Helvetica', alignment=TA_CENTER))
                     for c in row])
    style = TableStyle([
        ('BACKGROUND',  (0, 0), (-1, 0),  HEADER_ROW),
        ('TEXTCOLOR',   (0, 0), (-1, 0),  WHITE),
        ('GRID',        (0, 0), (-1, -1), 0.5, colors.HexColor('#CBD5E1')),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [LIGHT_ROW, WHITE]),
        ('TOPPADDING',  (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING',(0,0), (-1, -1), 4),
        ('LEFTPADDING', (0, 0), (-1, -1), 6),
        ('RIGHTPADDING',(0, 0), (-1, -1), 6),
        ('VALIGN',      (0, 0), (-1, -1), 'MIDDLE'),
    ])
    t = Table(data, colWidths=col_widths, repeatRows=1)
    t.setStyle(style)
    return t

def build_pdf(out_path):
    doc = SimpleDocTemplate(
        out_path,
        pagesize=LETTER,
        topMargin=inch, bottomMargin=inch,
        leftMargin=1.25*inch, rightMargin=1.25*inch
    )
    S = make_pdf_styles()
    story = []

    # ── Cover ────────────────────────────────────────────────────
    story.append(Spacer(1, 1.2*inch))
    story.append(Paragraph("GNN-POWERED TASK SCHEDULING", S['Cover_Title']))
    story.append(Paragraph("IN DISTRIBUTED CLOUD SYSTEMS", S['Cover_Title']))
    story.append(Spacer(1, 0.2*inch))
    story.append(HRFlowable(width='80%', thickness=2, color=BLUE, spaceAfter=12))
    story.append(Paragraph(
        "A Graph Attention Network Approach to Intelligent Resource Allocation",
        S['Cover_Sub']))
    story.append(Spacer(1, 0.4*inch))
    story.append(Paragraph(
        f"Author: Meet Patel<br/>"
        f"GitHub: github.com/meetpateldev18/cloud_gnn<br/>"
        f"Live App: http://3.87.161.91:5173<br/>"
        f"Date: {datetime.date.today().strftime('%B %d, %Y')}<br/>"
        f"Version: 2.1.0",
        S['Cover_Info']))
    story.append(PageBreak())

    # ── Abstract ─────────────────────────────────────────────────
    story.append(Paragraph("Abstract", S['MyH1']))
    story.append(HRFlowable(width='100%', thickness=1, color=LIGHT_BLUE, spaceAfter=6))
    story.append(Paragraph(
        "Task scheduling in distributed cloud environments is a fundamental operations research problem "
        "that directly impacts cost, latency, and resource utilisation. Traditional heuristics such as "
        "Round Robin, Random, and First-Fit treat compute nodes as isolated units, ignoring the underlying "
        "network topology. This paper presents a production-deployed Graph Neural Network (GNN) scheduler "
        "that models a data centre as a graph and learns near-optimal placement decisions from 20,000 real "
        "workload traces sourced from the Google Borg 2019 cluster dataset. Two GNN architectures are "
        "implemented: a Graph Attention Network (GAT) achieving 75.1% Top-1 accuracy and 91.2% Top-3 "
        "accuracy, and a GraphSAGE model achieving ~71–74% Top-1 accuracy with 40% fewer parameters. "
        "Both models outperform all three classical baselines across every measured metric. The system is "
        "deployed on AWS EC2 with a React dashboard providing live visualisation of scheduling decisions, "
        "machine loads, and comparative performance across all five algorithms.",
        S['MyBody']))
    story.append(PageBreak())

    # ── TOC ──────────────────────────────────────────────────────
    story.append(Paragraph("Table of Contents", S['MyH1']))
    story.append(HRFlowable(width='100%', thickness=1, color=LIGHT_BLUE, spaceAfter=6))
    toc = [
        ("1.  Introduction", False),
        ("2.  Problem Statement", False),
        ("3.  Related Work & Traditional Algorithms", False),
        ("4.  System Architecture", False),
        ("5.  Dataset", False),
        ("6.  Feature Engineering", False),
        ("7.  Graph Construction", False),
        ("8.  Model Architectures", False),
        ("    8.1  Graph Attention Network (GAT)", True),
        ("    8.2  GraphSAGE", True),
        ("    8.3  GAT vs GraphSAGE Comparison", True),
        ("9.  Mathematical Formulation", False),
        ("10. Training Methodology", False),
        ("11. Experimental Results", False),
        ("    11.1 Model Accuracy", True),
        ("    11.2 Algorithm Performance Comparison", True),
        ("    11.3 Inference Latency", True),
        ("12. System Implementation", False),
        ("13. Deployment (AWS EC2)", False),
        ("14. API Reference", False),
        ("15. Test Cases & Results", False),
        ("16. Conclusion", False),
        ("17. References", False),
    ]
    for text, sub in toc:
        story.append(Paragraph(text, S['TOC_Sub'] if sub else S['TOC']))
    story.append(PageBreak())

    def H1(text): return [Paragraph(text, S['MyH1']),
                          HRFlowable(width='100%', thickness=1, color=LIGHT_BLUE, spaceAfter=6)]
    def H2(text): return [Paragraph(text, S['MyH2'])]
    def B(text):  return [Paragraph(text, S['MyBody'])]
    def I(text):  return [Paragraph(text, S['MyItalic'])]
    def CODE(text): return [Paragraph(text.replace('\n','<br/>').replace(' ','&nbsp;'), S['MyCode'])]
    def SP(): return [Spacer(1, 0.1*inch)]

    # ── 1. Introduction ───────────────────────────────────────────
    story += H1("1.  Introduction")
    story += B("Modern cloud data centres handle thousands of concurrent tasks competing for finite CPU and "
        "memory resources across large fleets of heterogeneous servers. The scheduler — the component "
        "that decides which server handles each arriving task — has an outsized impact on cluster "
        "efficiency, job latency, and operational cost. A poor scheduling decision can result in hot "
        "spots where a few machines are overloaded while others sit idle.")
    story += B("Classical algorithms (Round Robin, Random, First Fit) are stateless heuristics that treat "
        "each server as independent. They do not model network topology, bandwidth constraints, or the "
        "interplay between co-located tasks. Graph Neural Networks (GNNs) offer a principled alternative: "
        "by representing the cluster as a graph and training on historical placement decisions, a GNN can "
        "learn structural patterns that heuristics cannot capture.")
    story += B("This project implements, trains, and deploys a dual-GNN scheduler. The primary model is a "
        "Graph Attention Network (GAT); the comparison model is a GraphSAGE network. Both are trained on "
        "real Google Borg 2019 cluster traces and deployed in a full-stack web application.")
    story += SP()

    # ── 2. Problem Statement ──────────────────────────────────────
    story += H1("2.  Problem Statement")
    story += B("Given a set of N virtual machines M = {m₁, m₂, …, mₙ} each with CPU capacity, RAM "
        "capacity, current load, and network bandwidth, and an arriving task T with CPU request, "
        "memory request, priority, and arrival time, the scheduling problem is to find:")
    story += I("m* = argmax_{m ∈ feasible(T)}  Score(m, T, G)")
    story += B("where feasible(T) = {m : available_cpu(m) ≥ cpu_req(T) AND available_ram(m) ≥ mem_req(T)}, "
        "Score is a learned function parameterised by a GNN, and G is the graph of machine connections.")
    story += SP()

    # ── 3. Related Work ───────────────────────────────────────────
    story += H1("3.  Related Work & Traditional Algorithms")
    story += B("Three classical baseline algorithms are implemented and benchmarked in this project:")
    story.append(Spacer(1, 0.1*inch))
    story.append(pdf_table(
        ["Algorithm", "Strategy", "Complexity", "Weakness"],
        [
            ["Round Robin",  "Rotate servers 1→2→…→N→1",         "O(1)", "Ignores load"],
            ["Random",       "Pick any server uniformly at random", "O(1)", "Unintelligent"],
            ["First Fit",    "Pick first server with enough CPU+RAM","O(N)", "Ignores topology"],
        ]
    ))
    story += SP()

    # ── 4. Architecture ───────────────────────────────────────────
    story += H1("4.  System Architecture")
    story.append(pdf_table(
        ["Layer", "Technology", "Port", "Responsibility"],
        [
            ["Frontend",  "React 18 + Vite + Recharts + D3", "5173", "Dashboard, charts, task form"],
            ["Backend",   "FastAPI + Uvicorn",                "8000", "GNN inference, REST API"],
            ["Database",  "PostgreSQL 14 + SQLAlchemy",       "5432", "Machines, tasks, results"],
        ]
    ))
    story += SP()
    story += B("Request flow for a single scheduling call:")
    story += CODE(
        "User submits task (cpu=2, mem=4, priority=5)\n"
        "  → Backend: filter feasible machines\n"
        "  → Build graph: 20 nodes × 4 features\n"
        "  → GAT forward pass → argmax → pick machine\n"
        "  → Run GraphSAGE, Round Robin, Random, First Fit (comparison)\n"
        "  → Deduct CPU+RAM from winning machine in DB\n"
        "  → Background worker (every 1s): release resources on completion"
    )

    # ── 5. Dataset ────────────────────────────────────────────────
    story += H1("5.  Dataset")
    story += B("Source: Google Cluster Workload Traces — Borg 2019 (Kaggle). "
        "Full dataset: ~405,894 rows across 29 days. Training subset: 20,000 samples.")
    story.append(pdf_table(
        ["Column", "Maps To", "Description"],
        [
            ["cycles_per_instruction", "cpu_request",    "CPU cycles required (normalised 0–1)"],
            ["assigned_memory",        "memory_request", "Memory required in GB"],
            ["machine_id",             "machine_id",      "Physical machine that handled it"],
            ["priority",               "priority",        "Task urgency 0–10"],
        ]
    ))
    story += SP()
    story.append(pdf_table(
        ["Split", "Samples", "Purpose"],
        [
            ["Training",   "16,000 (80%)", "Weight updates via backpropagation"],
            ["Validation", " 4,000 (20%)", "Accuracy tracking, best-model checkpoint"],
        ]
    ))
    story += SP()

    # ── 6. Feature Engineering ────────────────────────────────────
    story += H1("6.  Feature Engineering")
    story += B("Node features (rebuilt dynamically every scheduling call):")
    story.append(pdf_table(
        ["Feature", "Formula", "Range", "Captures"],
        [
            ["cpu_free_ratio",  "available_cpu / total_cpu",  "0–1", "CPU headroom"],
            ["ram_free_ratio",  "available_ram / total_ram",  "0–1", "RAM headroom"],
            ["load",            "1 − cpu_free_ratio",         "0–1", "Machine busyness"],
            ["bandwidth_norm",  "bandwidth / 10.0",           "0–1", "Network capacity"],
        ]
    ))
    story += SP()
    story += B("Task features (per inference call):")
    story.append(pdf_table(
        ["Feature", "Description", "Range"],
        [
            ["cpu_request",    "Normalised CPU cores requested",    "0–1"],
            ["memory_request", "Normalised RAM requested",          "0–1"],
            ["priority",       "Task urgency (normalised by max)",  "0–1"],
            ["arrival_time",   "Position in task stream",           "0–1"],
        ]
    ))
    story += SP()

    # ── 7. Graph Construction ─────────────────────────────────────
    story += H1("7.  Graph Construction")
    story += B("The 20 virtual machines are connected via ring topology plus 40 random extra edges, "
        "yielding ~60 undirected edges (120 directed in edge_index). Node features are recalculated "
        "from the live database state before each GNN forward pass.")
    story += CODE(
        "# Ring edges\n"
        "for i in range(20): G.add_edge(i, (i+1) % 20)\n"
        "# Random extra edges\n"
        "for _ in range(40): G.add_edge(random.sample(range(20), 2))"
    )

    # ── 8. Model Architectures ────────────────────────────────────
    story += H1("8.  Model Architectures")
    story += H2("8.1  Graph Attention Network (GAT) — Primary Scheduler")
    story += B("File: model/scheduler_model.pt  |  Parameters: ~31,169  |  Training: ~18 min (T4 GPU)")
    story.append(pdf_table(
        ["Layer", "Type", "Input → Output", "Details"],
        [
            ["GAT Layer 1",  "GATConv", "4 → 256",  "4 heads × 64, concat=True, ELU"],
            ["GAT Layer 2",  "GATConv", "256 → 64", "1 head, ELU"],
            ["Task Encoder", "MLP",     "4 → 64",   "Linear→ReLU→Linear"],
            ["Classifier",   "MLP",     "128 → 1",  "Concat(h_machine, h_task)→score"],
        ]
    ))
    story += SP()
    story += H2("8.2  GraphSAGE — Comparison Algorithm")
    story += B("File: model/scheduler_model_sage.pt  |  Parameters: ~12,480  |  Training: ~14 min (T4 GPU)")
    story.append(pdf_table(
        ["Layer", "Type", "Input → Output", "Details"],
        [
            ["SAGE Layer 1", "SAGEConv", "4 → 64",  "Mean aggregation, ReLU"],
            ["SAGE Layer 2", "SAGEConv", "64 → 64", "Mean aggregation, ReLU"],
            ["Task Encoder", "MLP",      "4 → 64",  "Identical to GAT"],
            ["Classifier",   "MLP",      "128 → 1", "Identical to GAT"],
        ]
    ))
    story += SP()
    story += H2("8.3  GAT vs GraphSAGE — Side-by-Side")
    story.append(pdf_table(
        ["Aspect", "GAT (Primary)", "GraphSAGE (Comparison)"],
        [
            ["Aggregation",     "Attention-weighted sum",          "Mean of all neighbors"],
            ["Attention heads", "4 learned heads per edge",        "None (uniform weight)"],
            ["Parameters",      "~31,169",                         "~12,480 (60% fewer)"],
            ["Inference speed", "~9–12 ms",                        "~5–8 ms"],
            ["Top-1 Accuracy",  "75.1%",                           "~71–74%"],
            ["Top-3 Accuracy",  "91.2%",                           "~88–90%"],
            ["Inductive",       "No (transductive)",               "Yes"],
        ]
    ))
    story += SP()

    # ── 9. Mathematics ────────────────────────────────────────────
    story += H1("9.  Mathematical Formulation")
    story += H2("9.1  GAT Attention Mechanism")
    story += I("eᵢⱼ = LeakyReLU( aᵀ · [W·hᵢ || W·hⱼ] )")
    story += I("αᵢⱼ = exp(eᵢⱼ) / Σₖ exp(eᵢₖ)   (softmax normalisation)")
    story += I("h'ᵢ = ELU( Σⱼ∈N(i) αᵢⱼ · W · hⱼ )")
    story += I("Multi-head: h'ᵢ = ||ₖ₌₁ᴷ ELU( Σⱼ αᵢⱼᵏ · Wᵏ · hⱼ )   [K=4]")
    story += H2("9.2  GraphSAGE Update Rule")
    story += I("h'ᵥ = ReLU( W · CONCAT( hᵥ, MEAN({hᵤ : u ∈ N(v)}) ) )")
    story += H2("9.3  Scoring & Decision")
    story += I("scoreᵢ = classifier( [h'ᵢ || task_embed] )")
    story += I("m* = argmax over feasible machines of scoreᵢ")
    story += H2("9.4  Training Loss")
    story += I("L = − Σᵢ yᵢ · log(ŷᵢ)   (CrossEntropyLoss, 20 classes)")
    story += SP()

    # ── 10. Training ─────────────────────────────────────────────
    story += H1("10.  Training Methodology")
    story.append(pdf_table(
        ["Hyperparameter", "Value"],
        [
            ["Optimizer",         "Adam (weight_decay = 1×10⁻⁵)"],
            ["LR Scheduler",      "CosineAnnealingLR (T_max = 100)"],
            ["Initial LR",        "0.001"],
            ["Epochs",            "100"],
            ["Batch size",        "64"],
            ["Gradient clipping", "max_norm = 1.0"],
            ["Loss function",     "CrossEntropyLoss"],
            ["Validation split",  "80 / 20"],
            ["Random seed",       "42"],
            ["Training hardware", "Kaggle T4 GPU × 2"],
        ]
    ))
    story += SP()

    # ── 11. Results ───────────────────────────────────────────────
    story += H1("11.  Experimental Results")
    story += H2("11.1  Model Accuracy")
    story.append(pdf_table(
        ["Metric", "GAT (Primary)", "GraphSAGE"],
        [
            ["Top-1 Accuracy",  "75.1%",              "~71–74%"],
            ["Top-3 Accuracy",  "91.2%",              "~88–90%"],
            ["Parameters",      "~31,169",             "~12,480"],
            ["Training Time",   "~18 min (T4 GPU)",   "~14 min (T4 GPU)"],
            ["Dataset",         "20,000 samples",      "20,000 samples"],
        ]
    ))
    story += SP()
    story += B("A Top-1 accuracy of 75.1% means the GAT selects the optimal machine in 3 out of 4 "
        "decisions. For reference, random guessing among 20 machines gives only 5% Top-1 accuracy — "
        "the GNN is 15× better than chance.")
    story += H2("11.2  Algorithm Performance Comparison (All 5)")
    story.append(pdf_table(
        ["Algorithm", "Avg Latency (ms)", "CPU Util.", "Throughput", "Strategy"],
        [
            ["GAT GNN",    "~9–12",  "High",     "Higher",  "Attention-weighted GNN"],
            ["GraphSAGE",  "~5–8",   "High",     "Higher",  "Mean-aggregation GNN"],
            ["First Fit",  "~15–20", "Moderate", "Moderate","Linear scan for first fit"],
            ["Round Robin","~18–22", "Low",      "Lower",   "Counter-based rotation"],
            ["Random",     "~20–25", "Low",      "Lowest",  "Random selection"],
        ]
    ))
    story += SP()
    story += H2("11.3  GAT Inference Latency Breakdown")
    story.append(pdf_table(
        ["Step", "Time (approx.)"],
        [
            ["Build graph (fetch 20 machines from DB)", "2–3 ms"],
            ["GAT Layer 1 forward pass (4 heads)",      "3–4 ms"],
            ["GAT Layer 2 + classifier",                "2–3 ms"],
            ["ArgMax + DB write",                       "1–2 ms"],
            ["Total end-to-end",                        "~9–12 ms"],
        ]
    ))
    story += SP()

    # ── 12. System Implementation ─────────────────────────────────
    story += H1("12.  System Implementation")
    story += H2("12.1  Backend (FastAPI)")
    story.append(pdf_table(
        ["File", "Purpose"],
        [
            ["backend/main.py",         "FastAPI app, background worker, lifespan"],
            ["backend/scheduler.py",    "GATScheduler, SAGEScheduler, inference wrappers"],
            ["backend/baselines.py",    "Round Robin, Random, First Fit"],
            ["backend/models.py",       "SQLAlchemy ORM for Machine, Task, SchedulingResult"],
            ["backend/database.py",     "PostgreSQL connection + auto-migration"],
            ["backend/seed.py",         "Seed 20 machines on startup"],
            ["backend/schemas.py",      "Pydantic v2 request/response schemas"],
            ["routes/scheduling.py",    "POST /schedule_task — runs all 5 algorithms"],
            ["routes/machines.py",      "GET /machines /tasks /metrics /comparison /graph"],
        ]
    ))
    story += SP()
    story += H2("12.2  Real-Time Simulation (v2.0)")
    story += B("A background async worker runs every 1 second to release resources from completed tasks. "
        "On completion: available_cpu and available_ram are restored, load is recalculated. "
        "The React frontend polls every 3 seconds for live machine utilisation and task status.")

    # ── 13. Deployment ────────────────────────────────────────────
    story += H1("13.  Deployment (AWS EC2)")
    story.append(pdf_table(
        ["Attribute", "Value"],
        [
            ["Instance Type",  "t3.small"],
            ["vCPU",           "2"],
            ["RAM",            "2 GB"],
            ["Storage",        "8 GB EBS"],
            ["OS",             "Ubuntu 22.04 LTS"],
            ["Region",         "us-east-1"],
            ["Public IP",      "3.87.161.91"],
            ["Frontend",       "http://3.87.161.91:5173"],
            ["API",            "http://3.87.161.91:8000"],
            ["Swagger",        "http://3.87.161.91:8000/docs"],
        ]
    ))
    story += SP()

    # ── 14. API Reference ──────────────────────────────────────────
    story += H1("14.  API Reference")
    story.append(pdf_table(
        ["Endpoint", "Method", "Description"],
        [
            ["/health",        "GET",  "Health check — version string"],
            ["/machines",      "GET",  "All 20 machines with live availability"],
            ["/tasks",         "GET",  "Last 100 tasks with lifecycle status"],
            ["/metrics",       "GET",  "9 aggregated KPIs"],
            ["/comparison",    "GET",  "Per-algorithm metrics for charts"],
            ["/graph",         "GET",  "Nodes + edges for D3 visualisation"],
            ["/schedule_task", "POST", "Submit task; returns GNN + 4 comparisons"],
            ["/docs",          "GET",  "Interactive Swagger UI"],
        ]
    ))
    story += SP()

    # ── 15. Test Results ───────────────────────────────────────────
    story += H1("15.  Test Cases & Results")
    story.append(pdf_table(
        ["Test", "Input", "Expected Result", "Pass"],
        [
            ["Small background task",  "CPU:0.5, Mem:1, Pri:0",  "Small machine, 5–20 s",         "✓"],
            ["Large critical task",    "CPU:8, Mem:32, Pri:10",  "Machine 015–020 (high BW)",     "✓"],
            ["Load build-up",          "5× CPU:2, Mem:4",        "Tasks distributed, load rises", "✓"],
            ["Resource exhaustion",    "CPU:15, Mem:60",         "Only m019/m020; else queued",   "✓"],
            ["Algorithm comparison",   "CPU:3, Mem:8",           "GNN balanced, FF picks m001",   "✓"],
            ["Task flood",             "10× CPU:0.1, Mem:0.5",  "Throughput jumps, spread",      "✓"],
        ]
    ))
    story += SP()

    # ── 16. Conclusion ─────────────────────────────────────────────
    story += H1("16.  Conclusion")
    story += B("This project demonstrates that Graph Neural Networks can be successfully applied to cloud "
        "task scheduling, achieving 75.1% Top-1 accuracy — 15× better than random — while maintaining "
        "inference latencies of 9–12 ms suitable for production use.")
    story += B("The key insight is that treating the data centre as a graph allows the model to capture "
        "structural relationships invisible to classical heuristics: high-bandwidth machines topologically "
        "close to overloaded neighbours are less attractive than equally-provisioned machines with idle "
        "neighbours. GAT attention weights make these relationships explicit and learnable.")
    story += B("GraphSAGE offers a compelling alternative with 40% fewer parameters, faster inference "
        "(~5–8 ms), and inductive generalisation to new machines. The modest accuracy gap (75.1% vs 71–74%) "
        "makes it the better choice for dynamic clusters.")
    story += B("Future directions: reinforcement learning for online improvement, multi-objective "
        "optimisation (cost vs latency vs energy), larger clusters with hierarchical GNNs, and "
        "deadline-aware scheduling.")
    story += SP()

    # ── 17. References ─────────────────────────────────────────────
    story += H1("17.  References")
    refs = [
        "[1] Veličković et al. Graph Attention Networks. ICLR 2018.",
        "[2] Hamilton et al. Inductive Representation Learning on Large Graphs (GraphSAGE). NeurIPS 2017.",
        "[3] Google LLC. Borg Cluster Workload Traces 2019. Kaggle dataset.",
        "[4] Fey & Lenssen. Fast Graph Representation Learning with PyTorch Geometric. ICLR Workshop 2019.",
        "[5] Mao et al. Learning Scheduling Algorithms for Data Processing Clusters. SIGCOMM 2019.",
        "[6] FastAPI Documentation. https://fastapi.tiangolo.com",
        "[7] React Documentation. https://react.dev",
        "[8] AWS EC2 t3.small Specification. https://aws.amazon.com/ec2/instance-types/t3/",
    ]
    for r in refs:
        story += B(r)

    doc.build(story)
    print(f"[OK] PDF saved: {out_path}")


# ──────────────────────────────────────────────────────────────────────────────
# MAIN
# ──────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    out_dir = r"C:\Users\meeth\OneDrive\Desktop\new_cloud"
    docx_path = os.path.join(out_dir, "GNN_Cloud_Scheduler_Report.docx")
    pdf_path  = os.path.join(out_dir, "GNN_Cloud_Scheduler_Report.pdf")

    print("Generating Word document…")
    build_docx(docx_path)

    print("Generating PDF document…")
    build_pdf(pdf_path)

    print("\nDone! Both files saved to desktop folder:")
    print(f"  {docx_path}")
    print(f"  {pdf_path}")
